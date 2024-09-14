# pip install python-docx

from docx import Document

def create():
    doc1 = Document()
    doc1.add_heading('如何使用python创建和操作word', 0)
    doc1.save('./generate_data/16_word_function.docx')

def headingAndParagraph():
    doc1 = Document()
    # 增加标题
    doc1.add_heading('如何使用python创建和操作word', 0)
    doc1.add_heading('这是一级标题', level= 1)
    doc1.add_heading('这是二级标题', level= 2)
    doc1.add_heading('这是三级标题', level= 3)

    # 增加段落
    doc1.add_paragraph('这是新的段落')
    doc1.add_paragraph('这是新的段落内容', 'Title') # 这个等同于add_heading中的零级标题

    paragraph1 = doc1.add_paragraph('在苍茫的大海上，狂风卷集着乌云。在乌云和大海之间，海燕像黑色的闪电，在高傲地飞翔。')
    paragraph1.add_run('一会儿翅膀碰着波浪，一会儿箭一般地直冲向乌云，它叫喊着，──就在这鸟儿勇敢的叫喊声里，乌云听出了欢乐。')
    paragraph2 = doc1.add_paragraph(' 《老人与海》是海明威的代表作，也是一部象征性的小说。主人公桑提亚哥是一位老渔夫，他经过重重艰险，捕获了“一条不止一千五百磅重的大马林鱼”，但这条大马林鱼却被鲨鱼吃光了，桑提亚哥只拖回了一副鱼的骨架。')
    paragraph2.add_run('本作品被称为是影响历史的百部经典之一；美国历史上里程碑式的32本书之一 。是世界文学宝库中的珍品，也是海明威全部创作中的瑰宝。1953年5月4日，海明威以作品《老人与海》获普利策奖。')


    # 增加无序列表
    doc1.add_paragraph('哪个不是你的喜欢的人物：')
    doc1.add_paragraph('孔子', style = 'List Bullet')
    doc1.add_paragraph('孟子', style = 'List Bullet')
    doc1.add_paragraph('老子', style = 'List Bullet')
    doc1.add_paragraph('荀子', style = 'List Bullet')
    doc1.add_paragraph('韩非子', style = 'List Bullet')
    
    # 增加有序列表
    doc1.add_paragraph('哪个不是你的喜欢的人：')
    doc1.add_paragraph('沫子', style = 'List Number')
    doc1.add_paragraph('豚豚', style = 'List Number')
    doc1.add_paragraph('UU', style = 'List Number')
    doc1.add_paragraph('奶兔', style = 'List Number')
    doc1.add_paragraph('奶茶', style = 'List Number')

    # 增加引用
    doc1.add_paragraph('这是一个引用', style= 'Intense Quote')

    # 增加图片
    pic = doc1.add_picture('./base_data/background.jpg')
    pic_width = pic.width
    pic_height = pic.height
    pg_width = doc1.sections[0].page_width
    sc = (pg_width/10 - doc1.sections[0].left_margin/10*2) / (pic_width/10)
    pic.width = int(pic_width * sc)
    pic.height = int(pic_height * sc)

    # 增加表格
    table = doc1.add_table(rows=1,cols=3)
    cells = table.rows[0].cells
    cells[0].text = '编号'
    cells[1].text = '姓名'
    cells[2].text = '职业'

    data = [
        ['1', '关羽', '将军'],
        ['2', '诸葛亮', '军师'],
        ['3', '刘备', '主帅'],
    ]

    for i,x,z in data:
        cells = table.add_row().cells # 增加一行表格
        cells[0].text = i # 如果是enumerate方法，则不是iterator，那么就不能是数值，必须是字符串string
        cells[1].text = x
        cells[2].text = z



    # 保存word
    doc1.save('./generate_data/16_word_function_add_heading.docx')


if __name__ == "__main__":
    #create()
    headingAndParagraph()