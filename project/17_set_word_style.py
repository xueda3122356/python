from docx import Document
from docx.shared import Inches, Pt, RGBColor

def use_styles():
    from docx.oxml.ns import qn
    doc1 = Document()

    p1 = doc1.add_paragraph('这是段落1:\n')
    p1.add_run('这是内容1.1\n').bold = True
    p1.add_run('这是内容1.2\n').italic = True
    p1.add_run('这是内容1.3\n').font.size = Pt(26)
    p1.add_run('这是内容1.4\n').font.strike = True
    p1.add_run('这是内容1.5\n').font.shadow = True
    p1.add_run('这是内容1.6\n').font.color.rgb = RGBColor(255,130,71)

    # 英文设置字体：可以直接设置
    p1.add_run('Test Font Style, 1.7\n').font.name = '微软雅黑'

    # 中文设置字体: 需要通过元素设置区域，从而选择字体格式
    run = p1.add_run('测试字体格式 1.8\n')
    
    # 需要初始化字体格式, 否则无法调用
    run.font.name = ''

    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # align 对齐
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p2 = doc1.add_paragraph('这是段落2:\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 全部向左缩进
    doc1.add_paragraph('中秋假期,广州白云机场口岸出入境旅客超8.3万人次。这标志着2024年以来广州白云机场口岸出入境客流量正式突破“1000万”节点,较2023年全年总量增长19%，位居全国空港口岸第三位。:\n').paragraph_format.left_indent = Inches(0.5)
    # 首行缩进
    doc1.add_paragraph('中秋假期,广州白云机场口岸出入境旅客超8.3万人次。这标志着2024年以来广州白云机场口岸出入境客流量正式突破“1000万”节点,较2023年全年总量增长19%，位居全国空港口岸第三位。:\n').paragraph_format.first_line_indent = Inches(0.5)


    # 设置段落之间的间距
    doc1.add_paragraph('这是段落3:\n').paragraph_format.space_before = Pt(30)
    doc1.add_paragraph('这是段落4:\n').paragraph_format.space_after = Pt(30)
    doc1.add_paragraph('这是段落5:\n')

    # 设置段落中的行间距
    doc1.add_paragraph('中秋假期,广州白云机场口岸出入境旅客超8.3万人次。这标志着2024年以来广州白云机场口岸出入境客流量正式突破“1000万”节点,较2023年全年总量增长19%，位居全国空港口岸第三位。:\n').paragraph_format.line_spacing = 3

    doc1.save('./generate_data/17_sytle_setting.docx')

if __name__ == "__main__":
    use_styles()
