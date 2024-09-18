from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_doc(carPlate,year,month,day,hour,minute,type_info,money):
    doc1 = Document()

    title = doc1.add_paragraph()
    p = title.add_run('车辆违章处罚通知书')

    p.font.size = Pt(30)
    p.font.color.rgb = RGBColor(255,0,0)
    p.font.name = ''
    p._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    content = doc1.add_paragraph()
    p2 = content.add_run(f'车号 {carPlate}车于{year}年{month}月{day}日{hour}时{minute}分在营运过程中出现（{type_info}）现象，公司按照安全法规和公司相关制度规定决定对该车驾驶员处以{money}元罚款，要求你在今后的营运过程中严格按照相关法律法规运行。（注：罚款金额请在返程后立即到公司缴纳\n驾驶员保证:\n驾驶员签字:\n年 月 日')
    content.paragraph_format.first_line_indent = Inches(0.3)
    p2.font.name = ''
    p2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc1.save('./generate_data/18_batch_generate_word.docx')

if __name__ == "__main__":
    carPlate = '桂A05834'
    year = 2024
    month = 8
    day = 5
    hour = 14
    minute = 38
    type_info = '违章'
    money = 500
    create_doc(carPlate,year,month,day,hour,minute,type_info,money)