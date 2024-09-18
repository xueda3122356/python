import docx

doc1 = docx.Document('./generate_data/16_word_function_add_heading.docx')

for p in doc1.paragraphs:
    print(p.text)

for t in doc1.tables:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text)